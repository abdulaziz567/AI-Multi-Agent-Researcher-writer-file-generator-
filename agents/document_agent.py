"""
Document Agent
--------------
Job: Take the final article (title + Markdown-ish body) and export it
as a DOCX and/or PDF file, saved to an output folder, and return the
file path(s) that the user can download.
"""

import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from fpdf import FPDF


def _slugify(text: str) -> str:
    text = re.sub(r"[^a-zA-Z0-9\s-]", "", text).strip().lower()
    text = re.sub(r"[\s_-]+", "_", text)
    return text[:60] if text else "article"


def _parse_sections(body: str):
    """
    Very small Markdown parser: splits body into a list of
    {"type": "heading"|"paragraph", "text": str}
    """
    blocks = []
    for raw_line in body.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            blocks.append({"type": "heading", "text": line[3:].strip()})
        elif line.startswith("# "):
            blocks.append({"type": "heading", "text": line[2:].strip()})
        else:
            blocks.append({"type": "paragraph", "text": line})
    return blocks


class DocumentAgent:
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def _base_filename(self, title: str) -> str:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{_slugify(title)}_{stamp}"

    def generate_docx(self, title: str, body: str) -> str:
        print("[DocumentAgent] Generating DOCX...")
        doc = Document()

        heading = doc.add_heading(title, level=0)

        for block in _parse_sections(body):
            if block["type"] == "heading":
                doc.add_heading(block["text"], level=1)
            else:
                p = doc.add_paragraph(block["text"])
                p.style.font.size = Pt(11)

        filename = self._base_filename(title) + ".docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def generate_pdf(self, title: str, body: str) -> str:
        print("[DocumentAgent] Generating PDF...")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 10, title)
        pdf.ln(4)

        for block in _parse_sections(body):
            if block["type"] == "heading":
                pdf.set_font("Helvetica", "B", 14)
                pdf.ln(3)
                pdf.multi_cell(0, 8, block["text"])
                pdf.ln(1)
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.multi_cell(0, 7, block["text"])
                pdf.ln(2)

        filename = self._base_filename(title) + ".pdf"
        filepath = os.path.join(self.output_dir, filename)
        pdf.output(filepath)
        return filepath

    def run(self, title: str, body: str, formats: list[str]) -> dict:
        """
        formats: list containing "docx" and/or "pdf"
        Returns dict of {"docx": path_or_None, "pdf": path_or_None}
        """
        results = {"docx": None, "pdf": None}
        if "docx" in formats:
            results["docx"] = self.generate_docx(title, body)
        if "pdf" in formats:
            results["pdf"] = self.generate_pdf(title, body)
        return results
